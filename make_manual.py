#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""고객컨설팅 마스터과정 운영관리 앱 사용 메뉴얼 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ────────────────────────────────────────────────────────────────
# 헬퍼
# ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e) if level == 1 else \
                             RGBColor(0x15, 0x65, 0xc0) if level == 2 else \
                             RGBColor(0x01, 0x57, 0x9b)
    return h

def add_para(doc, text, bold=False, size=10, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_number(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_screenshot_box(doc, desc, height_cm=4.5):
    """스크린샷 자리 표시자 박스"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'F0F4FF')
    cell.width = Cm(15)
    # 테두리
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '1565C0')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)

    icon_run = p.add_run('📸  ')
    icon_run.font.size = Pt(14)

    desc_run = p.add_run(f'[ 스크린샷 ]\n{desc}')
    desc_run.font.size = Pt(9)
    desc_run.font.color.rgb = RGBColor(0x15, 0x65, 0xc0)
    desc_run.italic = True

    doc.add_paragraph()
    return tbl

def add_tip_box(doc, text, tip_type='tip'):
    colors = {'tip': ('E8F5E9','2E7D32','💡 TIP'),
              'warn': ('FFF8E1','F57F17','⚠️ 주의'),
              'info': ('E3F2FD','1565C0','ℹ️ 참고')}
    bg, fg, label = colors.get(tip_type, colors['tip'])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r1 = p.add_run(f'{label}  ')
    r1.bold = True
    r1.font.size = Pt(10)
    r = int(fg[0:2], 16); g = int(fg[2:4], 16); b = int(fg[4:6], 16)
    r1.font.color.rgb = RGBColor(r, g, b)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    doc.add_paragraph()
    return tbl

def add_table_header(doc, headers, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '1565C0')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        run.font.size = Pt(9.5)
        if widths:
            cell.width = Cm(widths[i])
    return tbl

def add_table_row(tbl, values, bg='FFFFFF'):
    row = tbl.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(9)
    return row

def page_break(doc):
    doc.add_page_break()

# ────────────────────────────────────────────────────────────────
# 문서 초기화
# ────────────────────────────────────────────────────────────────
doc = Document()

# 페이지 여백 설정
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ════════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('고객컨설팅 마스터과정\n운영관리 시스템')
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph()

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('사용자 메뉴얼 v1.00')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

doc.add_paragraph()
doc.add_paragraph()

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_info.add_run(
    '현대해상화재보험  |  고객컨설팅 마스터과정\n'
    '문의: 호남지역단 이승학 전임강사\n\n'
    '최초 작성: 2026-05-20\n'
    '문서 버전: v1.00 (build 20260520u)'
)
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_paragraph()

# 표지 구분선
tbl_cover = doc.add_table(rows=1, cols=1)
cell_cover = tbl_cover.cell(0, 0)
set_cell_bg(cell_cover, '1565C0')
p_div = cell_cover.paragraphs[0]
p_div.paragraph_format.space_before = Pt(2)
p_div.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_note = p_note.add_run('이 메뉴얼은 앱 소스코드 분석을 기반으로 제작되었습니다.\n실제 화면은 메뉴얼의 스크린샷 설명란을 참고해 앱에서 직접 확인하세요.')
r_note.font.size = Pt(9)
r_note.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
r_note.italic = True

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 목차
# ════════════════════════════════════════════════════════════════
add_heading(doc, '목  차', 1)
toc_items = [
    ('1장', '앱 개요 및 접속 방법'),
    ('2장', '기본 화면 구성'),
    ('3장', '필터 사용법'),
    ('4장', '교육생 등록'),
    ('5장', '교육생 관리 및 면담 일지'),
    ('6장', '실적진도 탭'),
    ('7장', '시상 계산기'),
    ('8장', '팀 배정'),
    ('9장', '미지정 교육생 관리'),
    ('10장', '대시보드'),
    ('11장', '관리자 기능'),
    ('12장', '데이터 백업 · 복원'),
    ('13장', '시상안 편집'),
    ('14장', '자주 묻는 질문 (FAQ)'),
    ('[부록]', '버전 변경 이력'),
]
for ch, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r_ch = p.add_run(f'{ch}  ')
    r_ch.bold = True
    r_ch.font.size = Pt(10)
    r_ch.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    r_t = p.add_run(title)
    r_t.font.size = Pt(10)

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 1장. 앱 개요
# ════════════════════════════════════════════════════════════════
add_heading(doc, '1장. 앱 개요 및 접속 방법', 1)

add_para(doc, '고객컨설팅 마스터과정 운영관리 시스템은 현대해상화재보험의 마스터과정 교육생을 효율적으로 관리하기 위해 제작된 웹 앱입니다. 인터넷 브라우저만 있으면 PC · 태블릿 · 스마트폰에서 별도 설치 없이 사용 가능합니다.', size=10)
doc.add_paragraph()

add_heading(doc, '1.1 주요 기능', 2)
features = [
    ('📋 교육생 등록·수정·삭제', '지역단/비전센터/지점/기수/사번/이름/연락처 등 기본 정보 관리'),
    ('📝 면담 일지', '차수별 면담 기록 · 시상 계산기 · 코칭 포인트 저장 · 인쇄'),
    ('📈 실적진도', '교육생별 기준/현재실적, 신장률, 순증 실시간 현황 및 순위'),
    ('🏆 시상 계산기', '희망목표 입력 시 아너스클럽/하이포인트/마스터과정 예상 시상금 자동 계산'),
    ('👥 팀 배정', '조편성 숫자 직접 입력 또는 랜덤 자동 배정 + 조별 현황'),
    ('⚙️ 실적관리 (관리자)', '총괄월별실적·실적진도현황 붙여넣기, 아너스목표·인품실적 일괄 업데이트'),
    ('🏅 시상안 편집', '기수·스텝별 개인순증/신장률/신장액 Top N 시상 설정'),
    ('📦 백업·복원', '전체 데이터 JSON 파일로 내보내기/불러오기'),
]
tbl_feat = add_table_header(doc, ['기능', '설명'], [4, 11.5])
for feat, desc in features:
    add_table_row(tbl_feat, [feat, desc])
doc.add_paragraph()

add_heading(doc, '1.2 접속 방법', 2)
add_number(doc, 'Chrome 또는 Edge 브라우저를 실행합니다.')
add_number(doc, '주소창에 앱 URL을 입력하고 Enter를 누릅니다.')
add_number(doc, 'Firebase 서버와 연결이 완료되면 우측 상단에 ', bold_prefix='')
add_tip_box(doc, '앱은 GitHub Pages로 배포됩니다. 관리자로부터 URL을 받아 접속하세요. 즐겨찾기에 저장해 두면 편리합니다.', 'info')

add_screenshot_box(doc, '앱 첫 화면 — 상단 헤더 "Hyundai Marine & Fire Insurance" 로고, 대시보드/교육생 관리/실적진도/통계/관리자 메뉴, 우측 오류신고·새로고침 버튼과 연결 상태 표시')

add_heading(doc, '1.3 연결 상태 표시', 2)
add_bullet(doc, '초록색 "실시간 연결": Firestore 실시간 동기화 정상')
add_bullet(doc, '노란색 "캐시 표시중": 오프라인이지만 이전 데이터 표시 중')
add_bullet(doc, '빨간색 "연결 오류": 인터넷 연결 또는 Firebase 문제 → 🔄 버튼 클릭')
add_tip_box(doc, '🔄 (새로고침) 버튼은 단순 새로고침이 아니라 캐시를 초기화하고 최신 버전을 강제 로드합니다. 앱이 이상하게 보일 때 눌러보세요.', 'tip')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 2장. 기본 화면 구성
# ════════════════════════════════════════════════════════════════
add_heading(doc, '2장. 기본 화면 구성', 1)

add_para(doc, '앱은 크게 세 영역으로 구성됩니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '전체 화면 레이아웃 — 상단 헤더(네비게이션), 좌측 사이드바(필터·교육생 목록), 우측 메인 영역(대시보드/패널)')

add_heading(doc, '2.1 상단 헤더', 2)
tbl_h = add_table_header(doc, ['요소', '설명'], [4, 11.5])
rows_h = [
    ('대시보드', '전체 KPI 카드 및 지역단별 순위 현황'),
    ('교육생 관리', '선택한 교육생의 면담 일지 및 이력'),
    ('실적진도', '기수·지역단별 실적 현황 및 시상 순위'),
    ('통계', '통계 예약 화면 (추후 확장)'),
    ('관리자', '시상안 편집, 데이터 관리, 백업/복원 등'),
    ('🚨 오류신고', '앱 오류 발생 시 관리자에게 신고'),
    ('🔄 새로고침', '캐시 초기화 후 앱 최신 버전 강제 로드'),
    ('연결 상태', 'Firebase 실시간 연결 상태 표시'),
]
for r in rows_h:
    add_table_row(tbl_h, r)
doc.add_paragraph()

add_heading(doc, '2.2 좌측 사이드바', 2)
add_bullet(doc, '교육생 등록 버튼과 내보내기(CSV) 버튼')
add_bullet(doc, '미지정 교육생 경고 버튼 (지역단/비전센터 미배정 교육생이 있을 때만 표시)')
add_bullet(doc, '전체 교육생 수 표시')
add_bullet(doc, '필터 영역 (지역단/비전센터/지점/이름·사번 검색/기수/스텝)')
add_bullet(doc, '지점별 교육생 명단 (클릭 → 면담 관리 화면으로 이동)')
doc.add_paragraph()

add_heading(doc, '2.3 메인 영역', 2)
add_bullet(doc, '상단 탭 메뉴(대시보드·교육생 관리·실적진도 등)에 따라 콘텐츠가 교체됩니다.')
add_bullet(doc, '지역단 필터 선택 시: KPI 카드가 슬라이딩 순위 현황판으로 전환됩니다.')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 3장. 필터 사용법
# ════════════════════════════════════════════════════════════════
add_heading(doc, '3장. 필터 사용법', 1)
add_para(doc, '사이드바 상단의 필터 영역은 앱 전체에 공통으로 적용됩니다. 올바른 필터 설정이 시상 계산기, 실적진도, 면담 관리 모두에 영향을 줍니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '사이드바 필터 영역 — 지역단 드롭다운(예: 호남지역단), 비전센터 드롭다운, 지점 드롭다운, 이름·사번 검색창, 과정 기수 선택, 과정 스텝 선택(Step1/Step2), 필터 초기화 버튼')

add_heading(doc, '3.1 필터 항목별 설명', 2)
tbl_f = add_table_header(doc, ['필터', '역할', '주의사항'], [3, 6.5, 6])
filter_rows = [
    ('지역단', '전체 데이터의 기준 지역단 선택 (필수)', '미선택 시 실적진도·시상안이 정상 작동하지 않을 수 있습니다'),
    ('비전센터', '선택한 지역단 내 비전센터만 표시', '지역단 선택 후 활성화됨'),
    ('지점', '선택한 비전센터 내 지점만 표시', '비전센터 선택 후 활성화됨'),
    ('이름·사번 검색', '이름 또는 사번으로 교육생 즉시 검색', '한 글자만 입력해도 실시간 필터'),
    ('과정 기수', '1기~6기 중 선택, 시상 계산기·실적진도에 반영', '시상 계산기는 이 기수 기준으로 시상안을 불러옵니다'),
    ('과정 스텝', 'Step1 또는 Step2 구분', '실적진도의 기준/현재실적 열이 스텝에 따라 달라집니다'),
    ('필터 초기화', '모든 필터를 기본값으로 되돌림', ''),
]
for r in filter_rows:
    add_table_row(tbl_f, r)
doc.add_paragraph()

add_heading(doc, '3.2 필터 설정 순서 (권장)', 2)
add_number(doc, '지역단 선택 (예: 호남지역단)')
add_number(doc, '기수 선택 (예: 1기)')
add_number(doc, '스텝 선택 (Step1 또는 Step2)')
add_number(doc, '필요 시 비전센터·지점 추가 선택')
add_tip_box(doc, '필터는 자동 저장됩니다. 앱을 다시 열어도 마지막 필터 상태가 유지됩니다.', 'info')

add_heading(doc, '3.3 지점별 교육생 명단 사용법', 2)
add_para(doc, '필터 아래에 비전센터 → 지점 → 교육생 순으로 펼침 목록이 표시됩니다.', size=10)
add_bullet(doc, '비전센터 이름 클릭: 해당 센터 접기/펼치기')
add_bullet(doc, '지점 이름 클릭: 해당 지점 소속 교육생 목록 접기/펼치기')
add_bullet(doc, '교육생 이름 클릭: 오른쪽 메인 영역에 면담 관리 화면이 열림')
add_bullet(doc, '📌 지정 버튼 (미지정 센터 옆): 해당 교육생들의 비전센터를 일괄 지정')

add_screenshot_box(doc, '지점별 교육생 명단 — "🏢 영등포비전센터" 헤더 아래 "🏢 영등포지점" 소제목, 교육생 이름 목록. 숫자 "5/3" 표시(전체5명/면담완료3명)')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 4장. 교육생 등록
# ════════════════════════════════════════════════════════════════
add_heading(doc, '4장. 교육생 등록', 1)
add_para(doc, '새 교육생을 시스템에 등록하거나 기존 교육생 정보를 수정합니다. 사번이 동일하면 자동으로 기존 데이터를 덮어씁니다(upsert 방식).', size=10)
doc.add_paragraph()

add_heading(doc, '4.1 단건 등록', 2)
add_para(doc, '한 명씩 직접 입력하는 방식입니다.', size=10, bold=True)
doc.add_paragraph()

add_number(doc, '사이드바 상단의 [교육생 등록] 버튼을 클릭합니다.')
add_number(doc, '"교육생 등록 / 수정" 팝업이 열립니다. [단건 입력] 탭이 기본 선택됩니다.')
add_number(doc, '아래 항목을 순서대로 입력합니다.')
doc.add_paragraph()

add_screenshot_box(doc, '교육생 등록 팝업 (단건 입력 탭) — 지역단·비전센터·지점 드롭다운, 마스터과정 기수 선택, 사번 입력란(필수), 교육생 이름, 연락처, 평균실적(원), 마스터목표(원), 아너스목표(원), 조편성 입력란')

tbl_reg = add_table_header(doc, ['입력 항목', '설명', '필수'], [4, 9, 2.5])
reg_fields = [
    ('지역단', 'ORG_DATA에 등록된 지역단 목록에서 선택', '✅'),
    ('비전센터', '지역단 선택 시 자동으로 해당 센터 목록 활성화', '✅'),
    ('지점', '비전센터 선택 시 자동으로 해당 지점 목록 활성화', ''),
    ('마스터과정 기수', '1기~6기 중 해당 기수 선택', ''),
    ('사번', '사원 번호 (저장 키 — 동일 사번이면 업데이트됨)', '✅'),
    ('교육생 이름', '한글 실명', ''),
    ('연락처', '010-0000-0000 형식', ''),
    ('평균실적 (원)', '최근 6개월 평균 인보험 환산실적 (원 단위)', ''),
    ('마스터목표 (원)', '"선택▼" 버튼으로 평균+5만~30만 또는 직접 입력', ''),
    ('아너스목표 (원)', '아너스클럽 목표 금액 (원 단위)', ''),
    ('조편성', '팀 번호 숫자 입력 (예: 1, 2, 3 …)', ''),
]
for r in reg_fields:
    add_table_row(tbl_reg, r)
doc.add_paragraph()

add_number(doc, '입력 완료 후 하단 [저장] 버튼을 클릭합니다.')
add_number(doc, '동일 사번이 이미 있는 경우 덮어쓰기 확인 팝업이 뜹니다. 기존 값과 새 값을 비교 후 [덮어쓰기] 또는 [취소]를 선택합니다.')

add_tip_box(doc, '마스터목표 "선택▼" 버튼을 누르면 "평균 +5만원", "+10만원", "+20만원", "+30만원", "직접 입력" 빠른 선택 메뉴가 나옵니다. 평균실적을 먼저 입력해야 정확한 값이 계산됩니다.', 'tip')

add_screenshot_box(doc, '마스터목표 선택 팝업 — "선택▼" 버튼 클릭 후 나타나는 드롭다운: "평균 +5만원", "평균 +10만원", "평균 +20만원", "평균 +30만원", "직접 입력" 5개 항목')

add_heading(doc, '4.2 기존 교육생 수정', 2)
add_bullet(doc, '방법 1: 좌측 사이드바에서 교육생 이름 클릭 → 메인 영역 우상단 [✏️ 수정] 버튼 클릭')
add_bullet(doc, '방법 2: 좌측 사이드바에서 교육생 이름을 길게 누름')
add_para(doc, '수정 폼에는 기존 값이 자동으로 채워집니다. 변경하고 싶은 항목만 수정 후 [저장]을 클릭합니다.', size=10)
doc.add_paragraph()

add_heading(doc, '4.3 일괄 등록 (붙여넣기)', 2)
add_para(doc, '엑셀/스프레드시트에서 여러 명의 데이터를 한 번에 붙여넣어 등록합니다.', size=10, bold=True)
doc.add_paragraph()

add_screenshot_box(doc, '교육생 등록 팝업 (붙여넣기 탭) — 기수 선택 드롭다운(필수, 주황색 테두리), 형식 A/B 설명, 대형 텍스트 입력란, "미리보기 확인 후 저장" 버튼')

add_number(doc, '[교육생 등록] 클릭 → [붙여넣기(여러명)] 탭 선택')
add_number(doc, '상단 기수 드롭다운에서 기수를 선택합니다 (형식 A 필수).')
add_number(doc, '엑셀에서 데이터를 복사(Ctrl+C)하고 입력란에 붙여넣기(Ctrl+V)합니다.')
add_number(doc, '[📋 미리보기 확인 후 저장] 버튼을 클릭해 파싱 결과를 확인합니다.')
add_number(doc, '미리보기에서 열 매핑이 올바른지 확인 후 [✅ 이대로 저장]을 클릭합니다.')
doc.add_paragraph()

add_para(doc, '◎ 지원 형식', size=10, bold=True)
tbl_fmt = add_table_header(doc, ['형식', '열 수', '열 순서'], [2, 2, 11.5])
fmt_rows = [
    ('형식 A (권장)', '22열', '지역단 | 비전센터 | 지점 | 사원번호 | 성명 | 위촉차월 | 육성리더 | 인보험 | 환산 | 육성소득 | [기준/현재/달성률/순증] | 인품건수 | 인품실적'),
    ('형식 B', '11열', '지역단 | 비전센터 | 지점 | 기수 | 사번 | 이름 | 연락처 | 평균실적(원) | 마스터목표(원) | 아너스목표(원) | 차월(선택)'),
]
for r in fmt_rows:
    add_table_row(tbl_fmt, r)
doc.add_paragraph()

add_tip_box(doc, '형식 A는 실적진도 현황 보고서에서 그대로 복사해 붙여넣으면 됩니다. 열 매핑 팝업이 뜨면 각 열이 올바르게 인식됐는지 확인하고 "이 매핑으로 저장 진행"을 클릭하세요.', 'tip')
add_tip_box(doc, '붙여넣기 데이터에 없는 사번의 경우 신규 교육생 팝업이 뜹니다. 이름과 기수, 연락처를 확인하고 저장하면 자동으로 현재 지역단이 설정됩니다.', 'info')

add_heading(doc, '4.4 시드 파일 불러오기', 2)
add_para(doc, '관리자가 미리 준비한 시드 파일을 선택해 붙여넣기 입력란을 자동으로 채울 수 있습니다.', size=10)
add_number(doc, '일괄 등록 탭에서 [저장된 시드 파일 불러오기] 클릭')
add_number(doc, '지역단/기수별로 분류된 파일 목록에서 원하는 파일 선택')
add_number(doc, '텍스트 입력란에 자동으로 데이터가 채워집니다.')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 5장. 교육생 관리 및 면담 일지
# ════════════════════════════════════════════════════════════════
add_heading(doc, '5장. 교육생 관리 및 면담 일지', 1)
add_para(doc, '교육생을 선택하면 오른쪽 메인 영역에 면담 관리 화면이 나타납니다. 면담관리 / 면담이력 / 출력 세 개의 서브탭으로 구성됩니다.', size=10)
doc.add_paragraph()

add_heading(doc, '5.1 교육생 선택', 2)
add_bullet(doc, '좌측 사이드바 지점별 교육생 명단에서 교육생 이름을 클릭합니다.')
add_bullet(doc, '우측 메인 영역이 "교육생 면담 관리" 패널로 전환됩니다.')
add_bullet(doc, '상단에 교육생 기본 정보(이름/지역단/비전센터/지점/연락처/사번/기수) 카드가 표시됩니다.')

add_screenshot_box(doc, '교육생 선택 후 화면 — 상단에 교육생 카드(이름, 기수, 지점, 연락처, 사번), 우측 상단에 [✏️ 수정] 버튼, 하단에 면담관리/면담이력/출력 서브탭')

add_heading(doc, '5.2 면담관리 탭', 2)
add_para(doc, '새 면담 기록을 입력하거나 기존 면담 내용을 수정합니다.', size=10, bold=True)
doc.add_paragraph()

add_screenshot_box(doc, '면담관리 탭 — 상단에 시상 계산기 (희망목표, 평균실적, 진도% 등 입력란), 하단에 면담일지 서식 (차수, 현재실적, 진도%, 팀, 활동계획, 활동현황, 예상지출, 계약현황1/2, 코칭포인트), 저장 버튼')

add_para(doc, '◎ 시상 계산기 영역', size=10, bold=True)
tbl_calc = add_table_header(doc, ['항목', '설명'], [5, 10.5])
calc_fields = [
    ('희망목표 (원)', '이번 달 목표 실적 입력. 평균실적보다 높게 설정 권장'),
    ('"선택▼" 버튼', '평균+5만/10만/20만/30만원 빠른 선택 또는 직접 입력'),
    ('평균실적 (원)', '등록 시 입력한 평균실적 자동 표시 (수정 가능)'),
    ('인보험 순월납 (원)', '인보험 평균 월납입액 — 마스터과정 ③ 시상 계산에 사용'),
    ('계산 결과', '① 아너스클럽 + ② 하이포인트 + ③ 마스터 시상금 합계 표시'),
    ('④ 예상 시상', '실적진도 데이터 기반 신장률·신장액 순위시상 + 개인순증시상 예상'),
]
for r in calc_fields:
    add_table_row(tbl_calc, r)
doc.add_paragraph()

add_screenshot_box(doc, '시상 계산기 결과 화면 — ① 아너스클럽 테이블(현재 등급 강조), ② 개인 순증시상(하이포인트), ③ 마스터과정 개인시상(순증 기준), ④ 예상 시상(신장률 1위→시상품명), 최종 예상 시상금 합계, 🖨️ 시상인쇄 버튼')

add_para(doc, '◎ 면담일지 입력 항목', size=10, bold=True)
doc.add_paragraph()
tbl_iv = add_table_header(doc, ['항목', '설명', '자동 계산'], [4, 8, 3.5])
iv_fields = [
    ('면담 차수', '1차/2차/3차 등 숫자 입력 → 상단 제목 자동 업데이트', ''),
    ('현재실적 (원)', '이번 차수 기준 현재 실적', ''),
    ('진도 (%)', '현재실적/기준실적×100 자동 계산', '✅ 자동'),
    ('조 (팀 번호)', '팀 배정 숫자 (교육생 등록의 조편성과 연동)', ''),
    ('활동계획 (건)', '이번 달 계획 건수', ''),
    ('활동현황 (건)', '현재 달성 건수', ''),
    ('예상지출 (원)', '월 예상 지출 금액', ''),
    ('예상계약 (원)', '예상 계약 실적 금액', ''),
    ('확정계약 (원)', '확정된 계약 금액', ''),
    ('핵심 코칭포인트', '면담 핵심 내용, 후속 조치, 다음 주 계획 자유 기록', ''),
]
for r in iv_fields:
    add_table_row(tbl_iv, r)
doc.add_paragraph()

add_number(doc, '모든 항목 입력 후 [💾 저장] 버튼을 클릭합니다.')
add_number(doc, '저장 성공 시 "저장 완료" 토스트 메시지가 화면 하단에 표시됩니다.')
add_number(doc, '저장 후 면담이력 탭에서 방금 저장한 기록을 확인할 수 있습니다.')
add_tip_box(doc, '현재실적 입력 후 진도(%) 칸이 자동으로 채워집니다. 기준실적이 등록되어 있어야 합니다.', 'tip')

add_heading(doc, '5.3 면담이력 탭', 2)
add_para(doc, '저장된 모든 면담 기록을 차수 순서대로 조회합니다.', size=10)
add_bullet(doc, '각 면담 카드 클릭 → 해당 차수 내용이 면담관리 탭으로 불러와져 수정 가능')
add_bullet(doc, '면담 건수가 많으면 스크롤 가능')
add_bullet(doc, '상단 면담이력 탭 옆 숫자 배지가 전체 면담 기록 수를 표시')

add_screenshot_box(doc, '면담이력 탭 — 카드 형태로 1차, 2차, 3차... 면담 요약 표시. 각 카드에 차수·날짜·현재실적·진도%·핵심 코칭 요약이 나타남. 클릭하면 해당 차수 상세 편집 가능')

add_heading(doc, '5.4 출력 탭', 2)
add_para(doc, '선택한 교육생의 면담 일지를 인쇄합니다.', size=10)
add_bullet(doc, '출력 탭 클릭 → 인쇄 미리보기 페이지 표시')
add_bullet(doc, '브라우저 인쇄 기능(Ctrl+P) 또는 미리보기 내 인쇄 버튼 사용')
add_bullet(doc, '각 면담 차수별 1페이지씩 자동으로 분리 출력')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 6장. 실적진도 탭
# ════════════════════════════════════════════════════════════════
add_heading(doc, '6장. 실적진도 탭', 1)
add_para(doc, '상단 네비게이션에서 [실적진도]를 클릭합니다. 좌측 필터에서 지역단을 선택해야 데이터가 표시됩니다.', size=10)
doc.add_paragraph()

add_heading(doc, '6.1 실적진도 홈 화면', 2)
add_para(doc, '지역단·기수·스텝을 선택하면 홈 화면에 교육생별 실적 현황표가 나타납니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '실적진도 홈 화면 — 상단 지역단 태그·기수 선택·스텝 선택, 교육생별 테이블(이름·기준실적·현재실적·순증·신장률·인품건수·인품실적·조번호), 하단 요약 통계')

add_heading(doc, '6.2 실적진도 현황표 항목', 2)
tbl_pg = add_table_header(doc, ['열', '설명'], [4, 11.5])
pg_cols = [
    ('지점', '교육생 소속 지점'),
    ('이름', '교육생 이름'),
    ('기준실적', '활성화 마스터과정 시작 전 기준 실적 (원 단위)'),
    ('현재실적', '현재 달성 실적 (원 단위)'),
    ('신장률', '현재실적 ÷ 기준실적 × 100 (%)'),
    ('순증', '현재실적 - 기준실적 (원 단위)'),
    ('인품건수', '인품계약 건수'),
    ('인품실적', '인품계약 실적 금액'),
    ('조', '팀 배정 번호'),
]
for r in pg_cols:
    add_table_row(tbl_pg, r)
doc.add_paragraph()

add_heading(doc, '6.3 순위 버튼 클릭', 2)
add_para(doc, '홈 화면 상단에 있는 순위 버튼을 클릭하면 전체 순위 목록 팝업이 열립니다.', size=10)
add_bullet(doc, '📈 신장률 순위: 신장률 기준 내림차순 정렬 + 시상 할당 표시')
add_bullet(doc, '💰 신장액 순위: 순증액 기준 내림차순 정렬 + 시상 할당 표시')
add_bullet(doc, '👥 조별 현황: 팀별 합산 신장률/신장액 + 시상 현황')
add_bullet(doc, '🏅 개인순증 시상: 개인순증 기준 시상 목록')

add_screenshot_box(doc, '신장률 순위 팝업 — 1위부터 전체 순위표. 각 행에 이름·신장률·순증·시상명 표시. 중복 시상 방지(bothEnabled) 시 더 큰 시상만 표시되고 다른 항목은 회색 처리')

add_heading(doc, '6.4 슬라이딩 대시보드 (지역단 선택 시)', 2)
add_para(doc, '필터에서 지역단을 선택하면 메인 화면 상단 KPI 카드가 슬라이딩 현황판으로 전환됩니다.', size=10)
add_bullet(doc, '자동으로 5초마다 신장률 순위 → 신장액 순위 → 조별 현황 순서로 전환')
add_bullet(doc, '카드 좌우 화살표로 수동 이동 가능')

add_heading(doc, '6.5 ⚙️ 실적관리 (데이터 입력)', 2)
add_para(doc, '실적진도 탭에서 [⚙️ 실적관리 열기] 버튼을 클릭하면 데이터 입력 화면이 오버레이로 열립니다. 또는 관리자 탭 → 실적관리에서도 접근 가능합니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '실적관리 오버레이 — 상단에 4개 탭: "총괄월별실적 붙여넣기", "실적진도현황 붙여넣기", "아너스목표 붙여넣기", "인품실적 붙여넣기". 현재 활성 탭의 텍스트 입력란과 저장 버튼 표시')

add_para(doc, '◎ 총괄월별실적 붙여넣기', size=10, bold=True)
add_bullet(doc, '형식: "사번  장기하이캡  실적" (탭/공백 구분, 단위: 천원)')
add_bullet(doc, '평균실적과 하이캡 정보를 일괄 업데이트합니다.')
add_bullet(doc, '[📥 총괄월별실적 저장] 클릭 → 즉시 저장')

add_para(doc, '◎ 실적진도현황 붙여넣기', size=10, bold=True)
add_bullet(doc, '실적진도 현황 보고서에서 데이터를 그대로 복사해 붙여넣습니다.')
add_bullet(doc, '지역단·비전센터·지점·사번·성명·위촉차월·기준실적·현재실적·인품 등 자동 인식')
add_bullet(doc, 'Step2 선택 시: pgBase2, pgCurrent2, pgIpumCount2, pgIpumAmt2 필드로 저장')
add_bullet(doc, '열 매핑 팝업에서 각 열이 올바르게 인식됐는지 확인 후 저장')
add_bullet(doc, '시스템에 없는 사번은 신규 교육생으로 자동 등록 제안')

add_tip_box(doc, '기수와 스텝을 올바르게 선택한 후 붙여넣기를 진행하세요. 잘못된 기수로 저장된 데이터는 수정이 어렵습니다.', 'warn')

add_para(doc, '◎ 아너스목표 붙여넣기', size=10, bold=True)
add_bullet(doc, '형식: "사번  아너스목표금액(원)" (탭 구분)')
add_bullet(doc, '교육생의 아너스목표 항목을 일괄 업데이트합니다.')

add_para(doc, '◎ 인품실적 붙여넣기', size=10, bold=True)
add_bullet(doc, '형식: "사번  인품건수  인품실적(원)" (탭 구분)')
add_bullet(doc, '붙여넣기 적용 후 [✨ 인품 저장] 클릭')

add_heading(doc, '6.6 팀 배정 탭', 2)
add_para(doc, '실적관리 오버레이 내 [팀 배정] 탭에서 조편성을 관리합니다. (자세한 내용은 8장 참조)', size=10)

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 7장. 시상 계산기
# ════════════════════════════════════════════════════════════════
add_heading(doc, '7장. 시상 계산기', 1)
add_para(doc, '시상 계산기는 면담관리 탭 상단에 내장되어 있습니다. 교육생의 희망목표를 입력하면 세 가지 시상 항목의 예상 금액을 즉시 계산합니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '시상 계산기 입력 영역 — "희망목표(원)" 입력란 + "선택▼" 버튼, "평균실적(원)" 입력란, "인보험 순월납(원)" 입력란, "📊 계산" 버튼')

add_heading(doc, '7.1 입력값 설명', 2)
tbl_ci = add_table_header(doc, ['입력 항목', '설명', '없을 때'], [4.5, 7, 4])
ci_rows = [
    ('희망목표 (원)', '이번 달 달성하려는 목표 실적 (원 단위)', '계산 불가'),
    ('평균실적 (원)', '등록된 평균실적 자동 적용. 수정 가능', '0으로 처리'),
    ('인보험 순월납 (원)', '인보험 평균 월납입액. ③마스터 시상에만 사용', '0으로 처리'),
]
for r in ci_rows:
    add_table_row(tbl_ci, r)
doc.add_paragraph()

add_heading(doc, '7.2 계산 결과 구조', 2)
add_para(doc, '계산 버튼 클릭 또는 값 입력 시 아래 세 항목의 예상 시상금이 계산됩니다.', size=10)
doc.add_paragraph()

add_para(doc, '① 아너스클럽 시상', size=10, bold=True)
add_para(doc, '희망목표 기준으로 아너스클럽 등급을 판정하고 시상금을 표시합니다.', size=10, indent=0.5)
tbl_hon = add_table_header(doc, ['등급', '기준', '시상금 (만원)'], [5, 3, 3.5])
honors_data = [
    ('서밋 (Summit)', '500만원↑', '500'),
    ('로얄 마스터 (Royal master)', '400만원↑', '400'),
    ('엘리트 마스터 (Elite master)', '300만원↑', '250'),
    ('마스터 (Master)', '200만원↑', '150'),
    ('프레스티지Ⅱ (Prestige Ⅱ)', '150만원↑', '100'),
    ('프레스티지Ⅰ (Prestige Ⅰ)', '100만원↑', '70'),
    ('프라임Ⅱ (Prime Ⅱ)', '70만원↑', '50'),
    ('프라임Ⅰ (Prime Ⅰ)', '50만원↑', '30'),
    ('프로 (Pro)', '30만원↑', '20'),
]
for r in honors_data:
    add_table_row(tbl_hon, r)
doc.add_paragraph()

add_para(doc, '② 개인 순증시상 (하이포인트)', size=10, bold=True)
add_para(doc, '기본 지급 5만원 + 순증 × 50% (월 최대 20만원, 3개월 최대 50만원)', size=10, indent=0.5)
doc.add_paragraph()

add_para(doc, '③ 고객컨설팅마스터 개인시상', size=10, bold=True)
tbl_master = add_table_header(doc, ['기준 (마스터과정 순증)', '지급 방식', '월 시상금'], [5, 4, 4])
master_data = [
    ('순증 50만원↑', '순증 × 150%', '예: 순증 60만→90만원'),
    ('순증 30만원↑', '순증 × 120%', '예: 순증 40만→48만원'),
    ('순증 20만원↑', '20만원 고정', '20만원'),
    ('순증 10만원↑', '10만원 고정', '10만원'),
    ('순증 5만원↑', '5만원 고정', '5만원'),
]
for r in master_data:
    add_table_row(tbl_master, r)
add_para(doc, '* 마스터과정 순증 = 희망목표 - 인보험 순월납', size=9, color=(120,120,120))
doc.add_paragraph()

add_heading(doc, '7.3 ④ 예상 시상 (활성화 시상안 기준)', 2)
add_para(doc, '실적진도 데이터와 시상안 설정이 있을 때 표시됩니다. 현재 교육생의 실제 순위 기반으로 받게 될 시상을 카드 형태로 보여줍니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '④ 예상 시상 섹션 — "④ 예상 시상 | 1기 Step 1" 헤더. 시상 대상인 경우: 파란색 카드들 (📈 신장률 1위 / 137.2% → 다이슨V4 or 시디즈의자). 대상이 아닌 경우: 노란색 안내 (🎯 순증 N만원 추가 시 최소 개인순증시상 가능, 📈 신장률 N%p 추가 시 TopN 시상 진입 가능)')

add_bullet(doc, '시상 대상인 경우: 파란 카드로 시상 내용 표시 (예: 📈 신장률 1위 / 137.2% → 다이슨V4)')
add_bullet(doc, '시상 대상이 아닌 경우: 노란 안내 박스로 최소 시상 진입까지 부족분 안내')
add_bullet(doc, '신장률·신장액 중복 시 더 높은 시상 1개만 표시 (dedup 로직 적용)')
add_tip_box(doc, '④ 예상 시상은 시상안이 설정된 기수·스텝에서만 표시됩니다. 시상안 편집은 13장을 참조하세요.', 'info')

add_heading(doc, '7.4 🖨️ 시상 인쇄', 2)
add_bullet(doc, '[🖨️ 시상인쇄] 버튼 클릭 → 인쇄 미리보기 팝업')
add_bullet(doc, '교육생 이름·등급·금액이 포함된 시상 확인서 형태로 인쇄됩니다.')
add_bullet(doc, '여러 교육생을 연속 출력하려면 대시보드 → [🏆 시상안 출력] 버튼 사용')

add_heading(doc, '7.5 다음 단계 안내', 2)
add_para(doc, '계산 결과 하단에 "🚀 다음 단계 달성 목표" 박스가 표시됩니다.', size=10)
add_bullet(doc, '현재 등급에서 한 단계 올라가기 위해 희망목표를 얼마나 더 올려야 하는지 안내')
add_bullet(doc, '최상위 단계 달성 시: "🏆 최상위 단계 달성!" 표시')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 8장. 팀 배정
# ════════════════════════════════════════════════════════════════
add_heading(doc, '8장. 팀 배정 (조편성)', 1)
add_para(doc, '교육생들을 팀(조)별로 배정합니다. 교육생 등록/수정 폼의 "조편성" 란과 실적진도 실적관리의 "팀 배정" 탭이 동일한 데이터를 공유합니다.', size=10)
doc.add_paragraph()

add_heading(doc, '8.1 개별 팀 배정 (등록 폼 사용)', 2)
add_number(doc, '교육생 수정 폼을 엽니다 (사이드바에서 교육생 클릭 → [✏️ 수정])')
add_number(doc, '"조편성" 입력란에 숫자를 입력합니다 (예: 1, 2, 3)')
add_number(doc, '[저장]을 클릭합니다.')
add_tip_box(doc, '면담 관리 탭의 "조(팀 번호)" 란에 입력해 저장해도 교육생 레코드에 반영됩니다.', 'info')

add_heading(doc, '8.2 일괄 팀 배정 (실적관리 오버레이)', 2)
add_number(doc, '실적진도 탭 → [⚙️ 실적관리 열기] 클릭 (또는 관리자 → 실적관리)')
add_number(doc, '오버레이 내 [팀 배정] 탭 선택')
add_number(doc, '교육생 목록 테이블의 "팀" 열에 숫자를 직접 입력합니다.')
add_number(doc, '입력 완료 후 [💾 팀 배정 저장] 클릭')

add_screenshot_box(doc, '팀 배정 탭 — 교육생 테이블 (지점·이름·기준실적·현재실적·신장률·팀 입력란). 각 행의 "팀" 칸에 숫자 입력. 우측 상단에 "팀 수" 입력란과 "🎲 자동 배정" 버튼, 좌측에 조별 현황 카드')

add_heading(doc, '8.3 자동 배정', 2)
add_number(doc, '"팀 수" 입력란에 총 팀 개수를 입력합니다 (기본 8팀, 최대 20팀).')
add_number(doc, '[🎲 자동 배정 (랜덤 고르게)] 버튼 클릭')
add_number(doc, '확인 팝업에서 [예]를 클릭하면 Fisher-Yates 셔플로 고르게 배정됩니다.')
add_number(doc, '자동 배정 결과가 마음에 들지 않으면 직접 숫자를 수정합니다.')
add_number(doc, '[💾 팀 배정 저장] 클릭으로 확정합니다.')

add_tip_box(doc, '자동 배정은 화면에만 반영되고 저장은 되지 않습니다. 반드시 [💾 팀 배정 저장] 버튼을 눌러야 Firestore에 저장됩니다.', 'warn')

add_heading(doc, '8.4 조별 현황 보기', 2)
add_para(doc, '팀 배정 탭에서 데이터가 있으면 자동으로 조별 현황 요약 카드가 표시됩니다.', size=10)
add_bullet(doc, '각 조의 인원수·평균 신장률·합계 순증·시상 현황 요약')
add_bullet(doc, '실적진도 홈에서도 [👥 조별 현황] 버튼으로 팀 순위 확인 가능')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 9장. 미지정 교육생 관리
# ════════════════════════════════════════════════════════════════
add_heading(doc, '9장. 미지정 교육생 관리', 1)
add_para(doc, '지역단 또는 비전센터가 미배정된 교육생이 있으면 사이드바 상단에 ⚠️ 경고 버튼이 나타납니다. 이 버튼을 클릭해 일괄 배정하거나 불필요한 데이터를 삭제할 수 있습니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '사이드바 상단 ⚠️ 경고 버튼 — 주황/빨간 배경으로 "⚠️ 미지정 교육생 확인 (N명)" 표시. 클릭하면 미지정 교육생 목록 팝업 오픈')

add_heading(doc, '9.1 미지정 교육생 팝업 구성', 2)
add_bullet(doc, '교육생 목록: 이름, 현재 지역단, 비전센터, 지점, 기수, 조편성 정보')
add_bullet(doc, '지역단 드롭다운: 각 행별로 지역단 선택 가능')
add_bullet(doc, '비전센터 드롭다운: 선택한 지역단에 맞는 센터 목록 자동 갱신')
add_bullet(doc, '기수 선택: 해당 교육생의 기수 지정 가능')
add_bullet(doc, '삭제 체크박스: 체크하면 해당 행이 빨간색으로 변하고 저장 시 영구 삭제')

add_screenshot_box(doc, '미지정 교육생 팝업 — 테이블 형태. 각 행에 이름, 지역단 드롭다운, 비전센터 드롭다운, 기수 드롭다운, 삭제 체크박스(빨간색). 저장 버튼, 취소 버튼')

add_heading(doc, '9.2 비전센터 일괄 배정', 2)
add_para(doc, '교육생 목록 패널에서 "(비전센터 미지정)" 섹션 옆 [📌 지정] 버튼을 클릭합니다.', size=10)
add_number(doc, '비전센터 선택 목록이 팝업됩니다.')
add_number(doc, '상단에 "🔄 지점 기준 자동 매칭 (N명)" 버튼이 있으면 클릭해 지점→센터 자동 매칭 가능')
add_number(doc, '특정 센터 선택 → 확인 팝업 → [예, 저장] 클릭')
add_number(doc, '해당 미지정 교육생 전원의 비전센터가 일괄 설정됩니다.')

add_tip_box(doc, '지점 정보가 ORG_DATA에 등록된 경우 "🔄 지점 기준 자동 매칭" 옵션이 나타납니다. 한 번에 여러 센터로 자동 분류됩니다.', 'tip')

add_heading(doc, '9.3 교육생 삭제', 2)
add_number(doc, '⚠️ 버튼 클릭 → 미지정 교육생 팝업 오픈')
add_number(doc, '삭제할 교육생의 기수 열 "삭제" 체크박스에 체크합니다.')
add_number(doc, '행 배경이 빨간색으로 변합니다 (삭제 예정 표시).')
add_number(doc, '[💾 저장] 클릭 → 체크된 교육생은 영구 삭제, 나머지는 저장됩니다.')

add_tip_box(doc, '삭제는 취소할 수 없습니다. 반드시 이름을 확인 후 체크하세요. 삭제된 교육생의 면담 기록도 함께 삭제됩니다.', 'warn')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 10장. 대시보드
# ════════════════════════════════════════════════════════════════
add_heading(doc, '10장. 대시보드', 1)
add_para(doc, '앱 접속 시 기본으로 표시되는 화면입니다. 전체 또는 특정 지역단의 현황을 한눈에 볼 수 있습니다.', size=10)
doc.add_paragraph()

add_heading(doc, '10.1 지역단 미선택 상태', 2)
add_para(doc, '필터에서 지역단을 선택하지 않으면 아래 4개의 KPI 카드가 표시됩니다.', size=10)
tbl_kpi = add_table_header(doc, ['KPI 카드', '내용'], [5, 10.5])
kpi_rows = [
    ('전체 교육생', '등록된 전체 교육생 수'),
    ('평균실적 합계', '모든 교육생 평균실적(base) 합계 건수'),
    ('마스터목표 합계', '모든 교육생 마스터목표 합계 건수'),
    ('아너스목표 합계', '모든 교육생 아너스목표 합계 건수 (강조 표시)'),
]
for r in kpi_rows:
    add_table_row(tbl_kpi, r)
doc.add_paragraph()

add_heading(doc, '10.2 지역단 선택 시 — 슬라이딩 순위 현황판', 2)
add_para(doc, '좌측 필터에서 지역단을 선택하면 KPI 카드가 다음 슬라이딩 현황판으로 전환됩니다.', size=10)
add_bullet(doc, '신장률 TOP 순위 카드')
add_bullet(doc, '신장액 TOP 순위 카드')
add_bullet(doc, '조별 현황 카드')
add_bullet(doc, '5초마다 자동 전환 / 좌우 화살표로 수동 이동')

add_screenshot_box(doc, '슬라이딩 현황판 — "📈 신장률 TOP 현황 (1기 Step1)" 카드에 1위~5위 교육생 이름·신장률·시상 표시. 하단에 페이지 도트와 좌우 화살표 버튼')

add_heading(doc, '10.3 시상안 일괄 출력', 2)
add_para(doc, '대시보드 우측 상단 [🏆 시상안 출력] 버튼을 클릭하면 현재 필터 범위의 교육생 시상 예상 답안지를 일괄 출력합니다.', size=10)
add_bullet(doc, '지역단·비전센터·지점 필터 적용 범위 내 교육생만 출력')
add_bullet(doc, '각 교육생 1페이지씩 A4 사이즈로 출력')

add_heading(doc, '10.4 내보내기 (CSV)', 2)
add_para(doc, '사이드바 상단 [내보내기] 버튼 클릭 → 현재 필터 조건에 해당하는 교육생 목록을 CSV 파일로 다운로드합니다.', size=10)

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 11장. 관리자 기능
# ════════════════════════════════════════════════════════════════
add_heading(doc, '11장. 관리자 기능', 1)
add_para(doc, '상단 네비게이션 [관리자] 탭에서 접근합니다. 데이터 관리·편집 기능이 집중되어 있습니다.', size=10)
doc.add_paragraph()

add_screenshot_box(doc, '관리자 탭 화면 — 여러 섹션으로 구성: 오류신고 모음, 실적관리, 시상안 편집, 데이터 백업, 데이터 복원, 기본값 설정, 데이터 정비, 위험 구역, 앱 정보')

add_heading(doc, '11.1 오류신고 모음', 2)
add_bullet(doc, '앱 내 🚨 오류신고 버튼으로 제출된 신고 목록 조회')
add_bullet(doc, '선택 삭제 기능으로 처리 완료된 오류신고 정리')

add_heading(doc, '11.2 실적관리', 2)
add_bullet(doc, '[⚙️ 실적관리 열기] 버튼: 실적진도 데이터 입력 오버레이 오픈 (6장 참조)')

add_heading(doc, '11.3 기본값 설정', 2)
add_bullet(doc, '앱 기본 지역단 및 기본 시상 설정을 변경합니다.')

add_heading(doc, '11.4 데이터 정비', 2)
tbl_dm = add_table_header(doc, ['버튼', '기능'], [5, 10.5])
dm_rows = [
    ('📌 기수 미설정 → 1기 일괄 저장', '기수(cohort) 값이 없는 교육생 전원을 1기로 설정'),
    ('🗑️ 교육생 삭제 (개별 선택)', '체크박스로 원하는 교육생만 선택해 삭제'),
    ('🗑️ 필터 대상 일괄 삭제', '현재 필터 조건에 해당하는 교육생 전원 삭제 (매우 주의)'),
]
for r in dm_rows:
    add_table_row(tbl_dm, r)
add_tip_box(doc, '"필터 대상 일괄 삭제"는 매우 위험합니다. 실수로 대량 데이터가 삭제될 수 있으니 반드시 필터 범위를 확인하세요. 삭제 전 반드시 백업을 받아두세요.', 'warn')
doc.add_paragraph()

add_heading(doc, '11.5 앱 정보', 2)
add_bullet(doc, '현재 앱 버전 및 빌드 정보 표시')
add_bullet(doc, '예: v1.00 (build 20260520u)')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 12장. 데이터 백업 및 복원
# ════════════════════════════════════════════════════════════════
add_heading(doc, '12장. 데이터 백업 · 복원', 1)
add_para(doc, '모든 교육생 데이터를 JSON 파일로 내보내거나 기존 파일에서 복원할 수 있습니다.', size=10)
doc.add_paragraph()

add_heading(doc, '12.1 데이터 백업', 2)
add_number(doc, '관리자 탭 → [📦 데이터 백업] 섹션에서 [📦 데이터 백업] 버튼 클릭')
add_number(doc, '백업 팝업이 열리고 현재 등록된 전체 교육생 수가 표시됩니다.')
add_number(doc, '[💾 JSON 파일로 저장] 클릭 → 파일이 자동 다운로드됩니다.')
add_number(doc, '파일명: students_backup_YYYY-MM-DD.json 형식으로 저장됩니다.')

add_screenshot_box(doc, '데이터 백업 팝업 — "전체 N명 교육생 데이터를 백업합니다" 안내, "JSON 파일로 저장" 버튼')

add_heading(doc, '12.2 데이터 복원', 2)
add_number(doc, '관리자 탭 → [📥 데이터 복원] 섹션에서 [📥 JSON 파일에서 복원] 버튼 클릭')
add_number(doc, '파일 선택 대화상자에서 이전에 백업한 JSON 파일을 선택합니다.')
add_number(doc, '복원 방식을 선택합니다:')
add_bullet(doc, '추가/업데이트: 기존 데이터를 유지하면서 파일의 데이터를 병합', level=1)
add_bullet(doc, '전체 덮어쓰기: 기존 데이터를 모두 삭제하고 파일 데이터로 교체', level=1)
add_number(doc, '확인 팝업에서 [예, 복원합니다] 클릭')
add_tip_box(doc, '전체 덮어쓰기를 선택하면 현재 모든 데이터가 삭제됩니다. 먼저 현재 데이터를 백업한 후 복원하세요.', 'warn')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 13장. 시상안 편집
# ════════════════════════════════════════════════════════════════
add_heading(doc, '13장. 시상안 편집', 1)
add_para(doc, '기수·지역단·스텝별로 활성화 시상안을 설정합니다. 설정된 시상안은 실적진도 순위 화면과 시상 계산기의 ④ 예상 시상에 반영됩니다.', size=10)
doc.add_paragraph()

add_heading(doc, '13.1 시상안 편집기 열기', 2)
add_bullet(doc, '관리자 탭 → [시상안 편집] 섹션 → [🏆 시상안 편집] 클릭')
add_bullet(doc, '또는 실적진도 탭에서 시상 관련 버튼 클릭')

add_screenshot_box(doc, '시상안 편집 팝업 — 상단에 "연도·지역단·기수·Step 선택" 드롭다운, 불러오기·기본값·저장·닫기 버튼. 본문에 ① 개인순증시상 설정, ② Top N 시상1 설정, ③ Top N 시상2 설정, ④ 조시상, ⑤ 수상 자격 조건, ⑥ 메모 섹션')

add_heading(doc, '13.2 시상안 선택 방법', 2)
add_number(doc, '연도 드롭다운에서 연도 선택 (예: 2026)')
add_number(doc, '지역단 드롭다운에서 지역단 선택')
add_number(doc, '기수 드롭다운에서 기수 선택')
add_number(doc, 'Step 드롭다운에서 Step1 또는 Step2 선택')
add_number(doc, '[📂 불러오기] 버튼 클릭 → 저장된 시상안 목록에서 선택')

add_heading(doc, '13.3 개인순증시상 설정', 2)
add_bullet(doc, '"사용" 토글: 개인순증시상 활성화/비활성화')
add_bullet(doc, '[＋ 추가] 버튼: 시상 구간 추가')
add_bullet(doc, '각 구간: 기준금액(만원), 지급방식(현금/물품/비율%), 지급금액 설정')

add_heading(doc, '13.4 Top N 시상 설정 (신장률/신장액)', 2)
add_bullet(doc, '시상1/시상2 각각 "률(%)" 또는 "금액(원)" 타입 선택')
add_bullet(doc, '"N명" 설정: 시상 대상 인원 수 (예: 5 → Top 5까지 시상)')
add_bullet(doc, '[＋ 추가] 버튼으로 순위별 시상 추가 (1위, 2위, 3위…)')
add_bullet(doc, '각 순위: 현금(만원) 또는 물품명 입력')
add_tip_box(doc, '신장률·신장액 두 시상이 모두 설정된 경우, 중복 수상 방지 로직이 자동 적용됩니다. 두 항목 모두 해당 시 더 큰 시상 1개만 지급됩니다.', 'info')

add_heading(doc, '13.5 수상 자격 조건 설정', 2)
add_bullet(doc, '특정 조건을 충족한 교육생만 시상 대상으로 설정 가능')
add_bullet(doc, '조건 유형: 환산실적, 하이캡배수, 월납인보험 중 선택')
add_bullet(doc, 'AND/OR 연산자로 복수 조건 설정')

add_heading(doc, '13.6 시상안 저장', 2)
add_number(doc, '모든 설정 완료 후 [💾 저장] 버튼 클릭')
add_number(doc, '저장 확인 팝업: [💾 저장] 또는 [저장없이 닫기]')
add_number(doc, '저장된 시상안은 localStorage에 키 형식 "AP:연도:지역단:기수:Step"으로 보관됩니다.')
add_tip_box(doc, '시상안은 브라우저 localStorage에 저장됩니다. 다른 기기에서는 보이지 않으며, 브라우저 캐시 삭제 시 초기화됩니다. 중요한 시상안은 설정 후 스크린샷이나 메모로 별도 보관하세요.', 'warn')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 14장. FAQ
# ════════════════════════════════════════════════════════════════
add_heading(doc, '14장. 자주 묻는 질문 (FAQ)', 1)

faqs = [
    ('Q. 앱이 느리거나 데이터가 표시되지 않아요.',
     '우측 상단 🔄 버튼을 눌러 캐시를 초기화하세요. 인터넷 연결 상태도 확인하세요. 연결 상태 배지가 "연결 오류"로 표시되면 인터넷 상태를 점검하세요.'),
    ('Q. 교육생을 등록했는데 목록에 나타나지 않아요.',
     '좌측 필터가 다른 지역단/비전센터로 설정되어 있을 수 있습니다. [필터 초기화] 버튼을 클릭하거나 이름/사번으로 검색해보세요.'),
    ('Q. 시상 계산기가 작동하지 않아요 (④ 예상 시상이 안 보여요).',
     '필터에서 지역단·기수·스텝이 모두 선택되어 있는지 확인하세요. 시상안이 설정되지 않으면 ④ 예상 시상은 표시되지 않습니다. 13장을 참조해 시상안을 먼저 설정하세요.'),
    ('Q. Step2 시상 기준이 Step1으로 표시돼요.',
     '실적진도 탭 상단의 Step 선택과 좌측 필터의 스텝 선택이 일치하는지 확인하세요. Step2용 시상안을 별도로 등록해야 합니다.'),
    ('Q. 미지정 교육생이 계속 생겨요.',
     '실적진도 현황 붙여넣기에서 신규 사번이 감지될 때 지역단/비전센터 정보가 없으면 미지정으로 등록됩니다. 붙여넣기 전 필터에서 올바른 지역단을 선택하면 자동으로 채워집니다. 생성된 미지정 교육생은 9장의 방법으로 처리하세요.'),
    ('Q. 데이터를 실수로 삭제했어요.',
     '미리 백업한 JSON 파일이 있다면 12장의 복원 방법으로 복구하세요. 백업이 없으면 복구가 어렵습니다. 중요한 작업 전에는 반드시 백업하는 습관을 들이세요.'),
    ('Q. 여러 기기에서 동시에 사용할 수 있나요?',
     '네, Firebase Firestore 실시간 동기화로 여러 기기에서 동시 사용 가능합니다. 단, 시상안 설정은 localStorage에 저장되어 기기별로 별도 설정이 필요합니다.'),
    ('Q. CSV 내보내기 파일이 깨져 보여요.',
     '엑셀에서 열 때 파일 열기 마법사에서 인코딩을 UTF-8로 선택하거나, 파일 확장자를 .csv로 변경 후 메모장으로 열어 확인하세요.'),
    ('Q. 붙여넣기 시 열 매핑 팝업이 잘못 인식돼요.',
     '팝업에서 드롭다운으로 각 열의 의미를 수동으로 지정할 수 있습니다. 필요 없는 열은 "무시"로 설정하세요.'),
    ('Q. 앱 버전은 어디서 확인하나요?',
     '관리자 탭 → 앱 정보 섹션에서 현재 버전과 빌드 날짜를 확인할 수 있습니다. 화면 우측 상단 헤더에도 버전 배지가 표시됩니다.'),
]

for q, a in faqs:
    p_q = doc.add_paragraph()
    r_q = p_q.add_run(q)
    r_q.bold = True
    r_q.font.size = Pt(10)
    r_q.font.color.rgb = RGBColor(0x01, 0x57, 0x9b)
    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Cm(0.5)
    r_a = p_a.add_run(f'A. {a}')
    r_a.font.size = Pt(9.5)
    doc.add_paragraph()

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 부록: 버전 변경 이력
# ════════════════════════════════════════════════════════════════
add_heading(doc, '[부록] 버전 변경 이력', 1)
add_para(doc, '앱 개발 시작부터 현재까지의 모든 변경사항을 기록합니다.', size=10)
doc.add_paragraph()

tbl_ver = add_table_header(doc, ['버전', '빌드', '변경 내용', '비고'], [1.5, 2.5, 10.5, 1.5])

ver_history = [
    ('1.00', '20260520u', '버전 표기 체계 변경 — v2.03에서 v1.00으로 재시작. 이후 0.01씩 증가', '현재'),
    # 구버전 (이전 v2.x 체계)
    ('2.03', '20260520t', '미지정 교육생 생성 원인 차단 — 실적진도 붙여넣기 신규 등록 시 region/center 누락이면 현재 필터 지역단·ORG_DATA branch→center 맵으로 자동 보완', ''),
    ('2.02', '20260520s', '④ 예상 시상 섹션 위치 이동 (③ 개인시상 아래 / 최종합계 위). 시상 미해당 시 최소 시상 진입까지 부족분 안내 (노란 박스: 순증 N만원 추가 / 신장률 N%p 추가 등)', ''),
    ('2.01', '20260520r', '④ 예상 시상 섹션 개선 — 전체 시상안 목록 대신 해당 교육생이 실제 받게 될 시상만 카드 형태로 표시 (신장률·신장액 dedup 적용)', ''),
    ('2.00', '20260520q', '시상 계산기 하단에 ④ 활성화 시상안 섹션 추가. 기수·스텝별 시상 구간/순위 시상 표시', ''),
    ('1.99', '20260520p', '슬라이딩 대시보드 현황판 복원 — v1.98에서 _hrAwardPlan 변수 삭제로 사라진 문제 수정', ''),
    ('1.98', '20260520o', 'Step2 시상안 조회 오류 수정 — 기수 "기" 접미사 정규화 (예: "1기" → "1") 누락으로 Step1 시상안이 표시되던 버그 수정', ''),
    ('1.97', '20260520n', '사번 없는 교육생 삭제 기능 추가 — subscribe 스냅샷에 _docId 포함, DataAPI.removeByDocId() 메서드 신설', ''),
    ('1.96', '20260520m', '미지정 교육생 삭제 방식 변경 — 팝업 확인 방식에서 기수 열 삭제 체크박스로 변경. 체크 시 행 빨간색 강조', ''),
    ('1.95', '20260520l', '미지정 교육생 모달 개선 — 교육생 이름 좌측 삭제 버튼 추가, 비전센터 우측 기수 선택 셀렉트 추가', ''),
    ('1.94', '20260520k', '비전센터 미지정 일괄 지정 팝업 상단에 "🔄 지점 기준 자동 매칭 (N명)" 버튼 추가 — ORG_DATA branch→center 맵 활용', ''),
    ('1.93', '20260520j', '비전센터 미지정 센터 피커 — 교육생 데이터 대신 ORG_DATA 사용. 1개만 표시되던 버그 수정', ''),
    ('1.92', '20260520i', '비전센터 피커에서 "직접 입력" 옵션 제거. ORG_DATA 등록 센터만 선택 가능하도록 변경', ''),
    # 초기 기능 (구버전 이전)
    ('초기', '—', '교육생 등록·수정·삭제 (단건/일괄)', ''),
    ('초기', '—', '면담 일지 기록, 차수별 저장·수정·인쇄', ''),
    ('초기', '—', '시상 계산기 (아너스클럽/하이포인트/마스터과정)', ''),
    ('초기', '—', '실적진도 현황표 (기준/현재/신장률/순증/인품)', ''),
    ('초기', '—', '실적진도 데이터 붙여넣기 저장 (총괄월별/실적진도현황/아너스목표/인품실적)', ''),
    ('초기', '—', '신장률/신장액 Top N 순위 팝업 및 조별 현황', ''),
    ('초기', '—', '팀 배정 (수동 입력 + 랜덤 자동 배정)', ''),
    ('초기', '—', '데이터 백업/복원 (JSON 내보내기/불러오기)', ''),
    ('초기', '—', '시상안 편집 (기수·스텝별 개인순증/Top N/조시상/자격조건)', ''),
    ('초기', '—', '슬라이딩 대시보드 현황판 (지역단 선택 시)', ''),
    ('초기', '—', '지점별 교육생 명단, 비전센터별 접기/펼치기', ''),
    ('초기', '—', '필터 자동 저장 (localStorage)', ''),
    ('초기', '—', '오류신고 기능, 데이터 정비(기수 일괄 설정/개별·일괄 삭제)', ''),
    ('초기', '—', '조편성(팀 번호) 등록 폼 필드 추가', ''),
    ('초기', '—', '미지정 교육생 경고 버튼 및 일괄 지정 기능', ''),
    ('초기', '—', 'Firebase Firestore 실시간 동기화 (onSnapshot)', ''),
    ('초기', '—', '모바일 반응형 레이아웃 지원', ''),
]
for r in ver_history:
    alt = 'F8F9FA' if ver_history.index(r) % 2 == 0 else 'FFFFFF'
    if r[0] == '1.00':
        add_table_row(tbl_ver, r, bg='E8F5E9')
    elif r[2].startswith('초기'):
        add_table_row(tbl_ver, r, bg='FFFDE7')
    else:
        add_table_row(tbl_ver, r, bg=alt)

doc.add_paragraph()
add_tip_box(doc, '버전 표기 규칙: v1.00부터 시작해 기능 추가/버그 수정 시 0.01씩 증가합니다 (예: 1.00 → 1.01 → 1.02). 하위 호환성이 깨지는 대규모 변경 시에는 앞 자리수를 올립니다.', 'info')

page_break(doc)

# ════════════════════════════════════════════════════════════════
# 후기
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run(
    '이 메뉴얼은 앱 소스코드 및 기능 분석을 바탕으로 작성되었습니다.\n'
    '앱 업데이트 시 메뉴얼도 함께 갱신됩니다.\n\n'
    '문의: 호남지역단 이승학 전임강사\n'
    '앱 버전: v1.00  |  메뉴얼 작성일: 2026-05-20'
)
r_end.font.size = Pt(9)
r_end.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r_end.italic = True

# ════════════════════════════════════════════════════════════════
# 저장
# ════════════════════════════════════════════════════════════════
out_path = '/home/user/Customer-Master_online/고객컨설팅_마스터과정_사용자메뉴얼_v1.00.docx'
doc.save(out_path)
print(f'✅ 메뉴얼 생성 완료: {out_path}')
